"""
Exports generated articles to a .docx Word document.
"""
import io
from docx import Document
from docx.shared import Pt


def build_docx(title: str, content: str) -> bytes:
    doc = Document()

    title_para = doc.add_heading(title, level=1)

    for para in content.split("\n"):
        stripped = para.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        else:
            p = doc.add_paragraph(stripped)
            p.style.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
